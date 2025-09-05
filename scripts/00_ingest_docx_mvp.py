# -*- coding: utf-8 -*-
import argparse, os, json
from tqdm import tqdm
from docx import Document

def read_docx_text(path):
    doc = Document(path)
    lines = []
    for p in doc.paragraphs:
        t = (p.text or '').strip()
        if t:
            lines.append(t)
    for tbl in doc.tables:
        for r in tbl.rows:
            cells = [ (c.text or '').strip() for c in r.cells ]
            if any(cells):
                lines.append(' | '.join(cells))
    text = '\n'.join(lines)
    return text

def chunk(text, max_chars=1100, overlap=120):
    chunks = []
    i = 0
    while i < len(text):
        chunk = text[i:i+max_chars]
        if chunk.strip():
            chunks.append(chunk)
        i += max_chars - overlap
    return chunks

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--input', required=True, help='Folder with .docx files')
    ap.add_argument('--out', required=True, help='Output folder')
    args = ap.parse_args()

    os.makedirs(args.out, exist_ok=True)
    out_jsonl = os.path.join(args.out, 'chunks.jsonl')
    meta_jsonl = os.path.join(args.out, 'meta.jsonl')

    with open(out_jsonl, 'w', encoding='utf-8') as fout, open(meta_jsonl, 'w', encoding='utf-8') as fmeta:
        for name in tqdm(sorted(os.listdir(args.input))):
            if not name.lower().endswith('.docx'):
                continue
            path = os.path.join(args.input, name)
            try:
                text = read_docx_text(path)
                chunks = chunk(text)
                doc_id = os.path.splitext(name)[0]
                for i, ch in enumerate(chunks, 1):
                    rec = {
                        'id': f'{doc_id}_sec{i}',
                        'doc_id': doc_id,
                        'text': ch,
                        'source_path': path
                    }
                    fout.write(json.dumps(rec, ensure_ascii=False) + '\n')
                    fmeta.write(json.dumps(rec, ensure_ascii=False) + '\n')
            except Exception as e:
                print(f'[WARN] Failed {path}: {e}')

    print('Wrote:', out_jsonl, 'and', meta_jsonl)

if __name__ == '__main__':
    main()
