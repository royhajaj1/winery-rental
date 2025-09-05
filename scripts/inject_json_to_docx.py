import json
import os
from docx import Document
import re

def load_values(json_file):
    with open(json_file, 'r', encoding='utf-8') as f:
        data = json.load(f)
    print("Loaded JSON data:", data)
    
    # Map JSON keys to template placeholders
    key_mapping = {
        'DATE_CURRENT': 'CURRENT_DATE',
        'NAME': 'NAME',
        'PARTICIPANTS': 'PARTICIPANTS',
        'DATE_DO': 'DO_DATE',
        'HOUR_START': 'START_HOUR',
        'HOUR_FINISH': 'FINISH_HOUR',
        'COST_RENT': 'RENT_COST',
        'FOOD_MORNING': 'MORNING_FOOD',
        'FOOD_MAIN': 'MAIN_FOOD',
        'TOUR_WINE': 'WINE_TOUR'
    }
    
    values = {key_mapping[k]: str(v) for k, v in data.items() if k in key_mapping}
    print("Mapped values:", values)
    return values

def inject_values(template_file, output_file, values):
    doc = Document(template_file)
    found_placeholders = set()
    replacements_made = 0

    def replace_in_paragraph(paragraph):
        nonlocal replacements_made
        print(f"\nChecking paragraph: '{paragraph.text}'")
        
        # Store original text for verification
        original_text = paragraph.text
        
        # Get all runs and their text
        runs = paragraph.runs
        print(f"Number of runs: {len(runs)}")
        for i, run in enumerate(runs):
            print(f"Run {i}: '{run.text}'")

        # Check each placeholder
        for key, value in values.items():
            placeholder = f'X_{key}_X'
            if placeholder in original_text:
                print(f"Found placeholder: {placeholder} -> {value}")
                found_placeholders.add(placeholder)
                
                # If placeholder is in a single run, replace it there
                for run in runs:
                    if placeholder in run.text:
                        print(f"Replacing in run: '{run.text}' -> ", end='')
                        run.text = run.text.replace(placeholder, value)
                        print(f"'{run.text}'")
                        replacements_made += 1
                        break
                else:
                    # If placeholder spans multiple runs, reconstruct the paragraph
                    print("Placeholder spans multiple runs, reconstructing...")
                    new_text = original_text.replace(placeholder, value)
                    if runs:
                        runs[0].text = new_text
                        for run in runs[1:]:
                            run.text = ''
                        replacements_made += 1
                    print(f"After reconstruction: '{paragraph.text}'")

    # Process main document
    print("\nProcessing main document...")
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)

    # Process tables
    print("\nProcessing tables...")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)

    print(f'\nFound placeholders: {found_placeholders}')
    print(f'Made {replacements_made} replacements')
    print('Values used:', values)
    
    try:
        doc.save(output_file)
        print(f'\nSuccessfully saved document as {output_file}')
    except Exception as e:
        print(f'Error saving document: {str(e)}')

if __name__ == '__main__':
    import argparse
    
    parser = argparse.ArgumentParser(description='Inject JSON values into a Word document template.')
    parser.add_argument('json_file', help='Path to the JSON file with values')
    parser.add_argument('template_file', help='Path to the Word document template')
    
    args = parser.parse_args()
    
    # Convert to absolute paths
    json_path = os.path.abspath(args.json_file)
    template_path = os.path.abspath(args.template_file)
    
    # Generate output filename based on JSON filename
    json_name = os.path.splitext(os.path.basename(json_path))[0]
    output_path = os.path.join(os.path.dirname(json_path), f"{json_name}_proposal.docx")
    
    print("\n=== Starting document injection process ===")
    print(f"JSON Path: {json_path}")
    print(f"Template Path: {template_path}")
    print(f"Output Path: {output_path}")
    
    if not os.path.exists(json_path):
        print(f"Error: {json_path} does not exist!")
        exit(1)
    if not os.path.exists(template_path):
        print(f"Error: {template_path} does not exist!")
        exit(1)
        
    values = load_values(json_path)
    inject_values(template_path, output_path, values)
    print("=== Process completed ===\n")
